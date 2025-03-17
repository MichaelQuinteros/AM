# -*- coding: utf-8 -*-
"""app.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1An3HydE1rGaRw3Uq6kMlk4GkgHfrFOlH
"""

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm
import io

# Función para extraer datos de los archivos Excel
def extract_data_from_excels(excel_files, distrito, provincia, departamento):
    data_by_file = {}
    for excel_file in excel_files:
        try:
            df = pd.read_excel(excel_file)
            required_cols = {'Distrito', 'Provincia', 'Departamento'}
            if not required_cols.issubset(df.columns):
                st.warning(f"{excel_file} no contiene todas las columnas requeridas {required_cols}")
                continue

            df_filtered = df[(df['Distrito'].str.lower() == distrito.lower()) &
                             (df['Provincia'].str.lower() == provincia.lower()) &
                             (df['Departamento'].str.lower() == departamento.lower())]
            data_by_file[excel_file] = df_filtered
        except Exception as e:
            st.error(f"Error procesando {excel_file}: {e}")
    return data_by_file

# Reemplazar texto en el documento
def replace_text_in_document(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

# Insertar una tabla en el documento
def insert_table(doc, data, title):
    doc.add_paragraph(f"\n\n**{title}:**")
    table = doc.add_table(rows=1, cols=len(data.columns) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nro."
    hdr_cells[0].width = Cm(0.83)

    for i, column_name in enumerate(data.columns):
        hdr_cells[i + 1].text = column_name

    for idx, (_, row) in enumerate(data.iterrows(), start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[0].width = Cm(0.83)
        for i, value in enumerate(row):
            row_cells[i + 1].text = str(value)

# Generar el documento final
def generate_final_docx(output_docx, plantilla_docx, distrito, provincia, departamento, tables_by_file):
    doc = Document(plantilla_docx)

    selected_columns = ['PobrezaMonetaria', 'CCPP', 'Poblacion', 'Hogares', 'Viviendas', 'UnidadTerritorial']
    data_dict = {col: 'N/A' for col in selected_columns}

    if "datos1.xlsx" in tables_by_file and not tables_by_file["datos1.xlsx"].empty:
        data_dict.update(tables_by_file["datos1.xlsx"][selected_columns].iloc[0].to_dict())

    focalizado, seleccionado, mejorando, social = "NO", "NO", "NO", "NO"
    total_centros_poblados = "0"
    centros_poblados_df, infraestructura_social_df = pd.DataFrame(), pd.DataFrame()

    if "datos2.xlsx" in tables_by_file:
        table_cp = tables_by_file["datos2.xlsx"]
        if not table_cp.empty and all(col in table_cp.columns for col in ['Provincia', 'Distrito', 'CentroPoblado']):
            table_cp_filtered = table_cp[table_cp['Distrito'].str.lower() == distrito.lower()]
            focalizado = "SI" if not table_cp_filtered.empty else "NO"
            total_centros_poblados = str(len(table_cp_filtered))
            centros_poblados_df = table_cp_filtered[['Provincia', 'Distrito', 'CentroPoblado']]

    if "datos3.xlsx" in tables_by_file and 'Distrito' in tables_by_file["datos3.xlsx"]:
        seleccionado = "SI" if distrito.lower() in tables_by_file["datos3.xlsx"]['Distrito'].str.lower().values else "NO"

    if "datos4.xlsx" in tables_by_file and 'Distrito' in tables_by_file["datos4.xlsx"]:
        mejorando = "SI" if distrito.lower() in tables_by_file["datos4.xlsx"]['Distrito'].str.lower().values else "NO"

    if "datos5.xlsx" in tables_by_file:
        table_soc = tables_by_file["datos5.xlsx"]
        if not table_soc.empty and all(col in table_soc.columns for col in ['Provincia', 'Distrito', 'CentroPoblado']):
            social = "SI"
            infraestructura_social_df = table_soc[['Provincia', 'Distrito', 'CentroPoblado']]

    replacements = {
        "{Distrito}": distrito,
        "{Provincia}": provincia,
        "{Departamento}": departamento,
        "{PobrezaMonetaria}": str(data_dict.get("PobrezaMonetaria", "N/A")),
        "{CCPP}": str(data_dict.get("CCPP", "N/A")),
        "{Poblacion}": str(data_dict.get("Poblacion", "N/A")),
        "{Hogares}": str(data_dict.get("Hogares", "N/A")),
        "{Viviendas}": str(data_dict.get("Viviendas", "N/A")),
        "{UnidadTerritorial}": str(data_dict.get("UnidadTerritorial", "N/A")),
        "{Focalizado}": focalizado,
        "{Seleccionado}": seleccionado,
        "{Mejorando}": mejorando,
        "{Social}": social,
        "{TotalCentrosPoblados}": total_centros_poblados
    }
    replace_text_in_document(doc, replacements)

    if focalizado == "SI" and not centros_poblados_df.empty:
        insert_table(doc, centros_poblados_df, "Lista de Población Objetivo 2025")

    if social == "SI" and not infraestructura_social_df.empty:
        insert_table(doc, infraestructura_social_df, "Lista de Infraestructura Social")

    doc.save(output_docx)

# Streamlit interface
def main():
    st.title("Generador de Informe de Ayuda Memoria")

    # Cargar datos de 'datos1.xlsx' para las opciones de distrito, provincia y departamento
    df_datos1 = pd.read_excel("datos1.xlsx")

    # Extraer las opciones únicas de departamento
    departamentos = df_datos1['Departamento'].unique()

    # Combobox para seleccionar departamento
    departamento = st.selectbox("Selecciona el Departamento", departamentos)

    # Filtrar provincias basadas en el departamento seleccionado
    provincias = df_datos1[df_datos1['Departamento'] == departamento]['Provincia'].unique()

    # Combobox para seleccionar provincia (dependiente del departamento)
    provincia = st.selectbox("Selecciona la Provincia", provincias)

    # Filtrar distritos basados en la provincia seleccionada
    distritos = df_datos1[df_datos1['Provincia'] == provincia]['Distrito'].unique()

    # Combobox para seleccionar distrito (dependiente de la provincia)
    distrito = st.selectbox("Selecciona el Distrito", distritos)

    # Botón para generar el archivo
    if st.button("Generar Informe"):
        # Extraer datos de los Excel
        excel_files = ["datos1.xlsx", "datos2.xlsx", "datos3.xlsx", "datos4.xlsx", "datos5.xlsx"]
        tables_by_file = extract_data_from_excels(excel_files, distrito, provincia, departamento)

        # Nombre del archivo de salida
        output_docx = f"AyudaMemoria_{distrito}.docx"
        plantilla_docx = "plantilla.docx"

        # Generar el documento
        generate_final_docx(output_docx, plantilla_docx, distrito, provincia, departamento, tables_by_file)

        # Descargar el archivo generado
        with open(output_docx, "rb") as file:
            st.download_button("Descargar Informe", file, file_name=output_docx, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Ejecutar la aplicación Streamlit
if __name__ == "__main__":
    main()